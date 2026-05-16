from __future__ import annotations

from pathlib import Path
from time import perf_counter
from typing import Any

from docx import Document

from .base_converter import BaseConverter, ConversionError, ConversionResult
from .converter_factory import ConverterFactory


def _clean_text(value: str | None) -> str:
    return (value or "").strip()


def _paragraph_to_markdown(paragraph) -> str:
    text = _clean_text(paragraph.text)
    if not text:
        return ""

    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    low = style_name.lower()
    if low == "title":
        return f"# {text}"
    if low.startswith("heading "):
        level = style_name.split(" ", 1)[1]
        if level.isdigit():
            return f"{'#' * min(6, int(level) + 1)} {text}"
    return text


def _table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_clean_text(cell.text).replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


@ConverterFactory.register
class DOCXConverter(BaseConverter):
    supported_extensions = (".docx", ".doc")

    def convert(self, file_path: str | Path, output_dir: str | Path) -> ConversionResult:
        source = Path(file_path)
        if not source.exists():
            raise ConversionError(f"DOCX not found: {source}")
        if source.suffix.lower() not in self.supported_extensions:
            raise ConversionError(f"Unsupported file type: {source.suffix or '<none>'}")

        target_dir = Path(output_dir)
        target_dir.mkdir(parents=True, exist_ok=True)
        target_path = target_dir / f"{source.stem}.md"

        started = perf_counter()
        try:
            document = Document(str(source))
        except Exception as exc:
            raise ConversionError(f"Failed to read DOCX: {source.name}") from exc

        title = source.stem.replace("_", " ").replace("-", " ").strip().title() or source.stem
        sections: list[str] = []
        paragraph_count = 0

        for paragraph in document.paragraphs:
            rendered = _paragraph_to_markdown(paragraph)
            if rendered:
                sections.append(rendered)
                paragraph_count += 1

        table_markdowns: list[str] = []
        table_count = 0
        for table in document.tables:
            rendered = _table_to_markdown(table)
            if rendered:
                table_markdowns.append(rendered)
                table_count += 1

        metadata: dict[str, Any] = {
            "source_file": source.name,
            "markdown_path": str(target_path),
            "paragraphs": paragraph_count,
            "tables": table_count,
            "conversion_method": "python-docx",
            "conversion_time": round(perf_counter() - started, 3),
            "title": title,
        }

        frontmatter = [
            "---",
            f"title: {title}",
            f"source_file: {source.name}",
            f"paragraphs: {paragraph_count}",
            f"tables: {table_count}",
            "conversion_method: python-docx",
            "---",
        ]

        body: list[str] = [f"# {title}"]
        if sections:
            body.extend(sections)
        if table_markdowns:
            body.extend(table_markdowns)
        if len(body) == 1:
            body.append("_No extractable content was found in this DOCX._")

        markdown = "\n\n".join(frontmatter + body) + "\n"
        target_path.write_text(markdown, encoding="utf-8")

        return ConversionResult(
            markdown_path=target_path,
            metadata=metadata,
            status="success",
            content=markdown,
        )

