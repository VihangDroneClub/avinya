from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from converters.base_converter import ConversionError
from converters.docx_converter import DOCXConverter
from converters.xlsx_converter import XLSXConverter


def _build_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Vihang Report", level=0)
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("This is the weekly status report.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Battery"
    table.cell(1, 1).text = "Ready"
    doc.save(path)


def _build_sample_xlsx(path: Path) -> None:
    workbook = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
          xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets>
    <sheet name="Budget" sheetId="1" r:id="rId1"/>
  </sheets>
</workbook>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>
"""
    sheet = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1">
      <c r="A1" t="inlineStr"><is><t>Item</t></is></c>
      <c r="B1" t="inlineStr"><is><t>Amount</t></is></c>
    </row>
    <row r="2">
      <c r="A2" t="inlineStr"><is><t>Battery</t></is></c>
      <c r="B2"><v>15000</v></c>
    </row>
  </sheetData>
</worksheet>
"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>
"""
    root_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>
"""

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", root_rels)
        zf.writestr("xl/workbook.xml", workbook)
        zf.writestr("xl/_rels/workbook.xml.rels", rels)
        zf.writestr("xl/worksheets/sheet1.xml", sheet)


def test_docx_converter_writes_markdown(tmp_path: Path):
    source = tmp_path / "meeting_notes.docx"
    _build_sample_docx(source)

    output_dir = tmp_path / "vault"
    result = DOCXConverter().convert(source, output_dir)

    assert result.status == "success"
    assert result.markdown_path == output_dir / "meeting_notes.md"
    assert result.markdown_path.exists()
    assert result.metadata["source_file"] == "meeting_notes.docx"
    assert result.metadata["paragraphs"] >= 2
    assert result.metadata["tables"] == 1
    assert "# Vihang Report" in result.content
    assert "## Overview" in result.content
    assert "| Item | Value |" in result.content


def test_docx_converter_rejects_missing_file(tmp_path: Path):
    with pytest.raises(ConversionError):
        DOCXConverter().convert(tmp_path / "missing.docx", tmp_path / "vault")


def test_xlsx_converter_writes_markdown(tmp_path: Path):
    source = tmp_path / "budget.xlsx"
    _build_sample_xlsx(source)

    output_dir = tmp_path / "vault"
    result = XLSXConverter().convert(source, output_dir)

    assert result.status == "success"
    assert result.markdown_path == output_dir / "budget.md"
    assert result.markdown_path.exists()
    assert result.metadata["source_file"] == "budget.xlsx"
    assert result.metadata["sheets"] == 1
    assert "## Sheet: Budget" in result.content
    assert "| Item | Amount |" in result.content
    assert "| Battery | 15000 |" in result.content


def test_xlsx_converter_rejects_missing_file(tmp_path: Path):
    with pytest.raises(ConversionError):
        XLSXConverter().convert(tmp_path / "missing.xlsx", tmp_path / "vault")

